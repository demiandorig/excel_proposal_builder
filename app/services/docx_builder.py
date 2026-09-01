"""
Builds the client-facing email Word document (.docx).
Falls back gracefully when python-docx is not installed.
"""
from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# Entravision brand purple as RGB
_PURPLE = (91, 45, 142)   # #5B2D8E


def build_client_email_docx(
    output_path: Path,
    proposal_title: str,
    ae_name: str,
    ae_email: str,
    subject: str,
    body: str,
) -> bool:
    """
    Write the client-facing email as a formatted Word document.
    Returns True on success, False if python-docx is unavailable.
    """
    if not _HAS_DOCX:
        return False

    doc = Document()

    # --- Document title (proposal convention name) ---
    h = doc.add_heading(proposal_title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*_PURPLE)

    # --- Email header block ---
    subj_para = doc.add_paragraph()
    subj_run = subj_para.add_run("Subject: ")
    subj_run.bold = True
    subj_para.add_run(subject)

    doc.add_paragraph("")  # blank line

    # --- Email body ---
    _render_body(doc, body)

    # --- Sign-off footer ---
    if ae_name or ae_email:
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        if ae_name:
            r = footer.add_run(ae_name)
            r.bold = True
        if ae_email:
            doc.add_paragraph(ae_email)

    doc.save(str(output_path))
    return True


def _render_body(doc: "Document", body: str) -> None:
    """
    Convert the AI-generated email body (plain text with light markdown)
    into formatted Word paragraphs.

    Supported patterns:
      - Lines starting with "- " or "• "  →  bullet list
      - Lines in ALL CAPS or ending with ":"  →  bold paragraph (section heading)
      - Blank lines  →  paragraph break
      - Everything else  →  normal paragraph
    """
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            # Blank line → spacer paragraph
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("• "):
            # Bullet item
            text = stripped.lstrip("-•").strip()
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Check for heading-like lines (e.g. "**What This Campaign Delivers:**")
        heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        is_heading = (
            stripped.startswith("**") and stripped.endswith("**")
            or (stripped.endswith(":") and len(stripped) < 80 and stripped == stripped.upper())
        )
        if is_heading:
            p = doc.add_paragraph()
            r = p.add_run(heading_text)
            r.bold = True
            i += 1
            continue

        # Normal paragraph — render inline **bold** markers
        p = doc.add_paragraph()
        _render_inline(p, stripped)
        i += 1


def _render_inline(para, text: str) -> None:
    """Add runs to `para`, converting **bold** markers to bold runs."""
    import re as _re
    parts = _re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        else:
            para.add_run(part)


def build_roadblocks_docx(
    output_path: Path,
    title: str,
    overall_summary: str,
    product_roadblocks: list,
    used_web_search: bool,
) -> bool:
    """
    Write the Step 05 AI Roadblocks / Restrictions Check as a Word document —
    one section per product with its risk level, risks (issue/detail/source),
    and recommended mitigation.
    Returns True on success, False if python-docx is unavailable.
    """
    if not _HAS_DOCX:
        return False

    doc = Document()

    h = doc.add_heading(f"{title} — Platform Restrictions & Roadblocks", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*_PURPLE)

    if not used_web_search:
        note = doc.add_paragraph()
        note_run = note.add_run(
            "⚠ Generated without live web search — verify against current "
            "platform policies before relying on this for launch decisions."
        )
        note_run.italic = True

    if overall_summary:
        doc.add_paragraph("")
        summary_label = doc.add_paragraph()
        summary_label.add_run("Overall Risk Summary").bold = True
        doc.add_paragraph(overall_summary)

    for item in product_roadblocks:
        doc.add_paragraph("")
        h2 = doc.add_heading(item.get("product_name", ""), level=2)
        for run in h2.runs:
            run.font.color.rgb = RGBColor(*_PURPLE)

        risk_level = (item.get("risk_level") or "low").upper()
        risk_para = doc.add_paragraph()
        risk_para.add_run("Risk Level: ").bold = True
        risk_para.add_run(risk_level)

        for risk in item.get("risks") or []:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(risk.get("issue", "")).bold = True
            if risk.get("detail"):
                p.add_run(" — " + risk["detail"])
            if risk.get("source"):
                src = doc.add_paragraph()
                src_run = src.add_run(f"Source: {risk['source']}")
                src_run.italic = True
                src.paragraph_format.left_indent = Pt(24)

        if item.get("recommended_mitigation"):
            mit = doc.add_paragraph()
            mit.add_run("Recommended mitigation: ").bold = True
            mit.add_run(item["recommended_mitigation"])

    doc.save(str(output_path))
    return True


def build_strategy_brief_docx(
    output_path: Path,
    title: str,
    client_summary: str,
    market_context: str,
    objectives_analysis: str,
    strategy_summary: str,
    recommended_tactics: list,
    key_insights: list,
    monthly_budget: float = 0.0,
    total_months: int = 0,
) -> bool:
    """
    Write the Step 03 AI Strategy Brief as a Word document, in the same
    Entravision-branded format as the roadblocks report and client email —
    client/business summary, market context, objectives analysis, the
    recommended tactic mix with rationale/data/budget split, the overall
    strategy direction, and key insights for the AE.
    Returns True on success, False if python-docx is unavailable.
    """
    if not _HAS_DOCX:
        return False

    doc = Document()

    h = doc.add_heading(f"{title} — AI Strategy Brief", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*_PURPLE)

    if monthly_budget and total_months:
        budget_para = doc.add_paragraph()
        budget_run = budget_para.add_run(
            f"Budget: ${monthly_budget:,.0f}/month × {total_months} months "
            f"= ${monthly_budget * total_months:,.0f} total flight"
        )
        budget_run.italic = True

    def _section(label: str, body: str) -> None:
        if not body:
            return
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        doc.add_paragraph(body)

    _section("Client & Business", client_summary)
    _section("Market Context", market_context)
    _section("Objectives Analysis", objectives_analysis)

    if recommended_tactics:
        doc.add_paragraph("")
        h2 = doc.add_heading("Recommended Media Tactics", level=2)
        for run in h2.runs:
            run.font.color.rgb = RGBColor(*_PURPLE)

        for t in recommended_tactics:
            p = doc.add_paragraph()
            head_run = p.add_run(t.get("product_family", ""))
            head_run.bold = True
            pct = t.get("suggested_budget_pct")
            if pct:
                p.add_run(f"  —  {pct}% of budget")

            if t.get("rationale"):
                doc.add_paragraph(t["rationale"])
            if t.get("data_point"):
                data_p = doc.add_paragraph(style="List Bullet")
                data_p.add_run(t["data_point"])
                if t.get("citation"):
                    cite_run = data_p.add_run(f" ({t['citation']})")
                    cite_run.italic = True
            if t.get("entravision_advantage"):
                adv_p = doc.add_paragraph(style="List Bullet")
                adv_p.add_run("Entravision advantage: ").bold = True
                adv_p.add_run(t["entravision_advantage"])
            doc.add_paragraph("")

    _section("Strategy Direction", strategy_summary)

    if key_insights:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Key Insights for the AE").bold = True
        for insight in key_insights:
            doc.add_paragraph(insight, style="List Bullet")

    doc.save(str(output_path))
    return True


# Make re available at module level for _render_body
import re
